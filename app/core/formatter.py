"""
排版引擎模块 - 根据分析结果生成符合公文格式的Word文档
"""
import re
from pathlib import Path
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from app.core.styles import (
    ElementType,
    GovDocFonts,
    GovDocParagraphs,
    PageSettings,
    FONT_STYLE_MAP,
    PARAGRAPH_STYLE_MAP,
    FontStyle,
    ParagraphStyle
)
from app.core.llm_analyzer import AnalysisResult, DocumentElement


# 括号标准化映射表 - 将各种括号统一转换为中文全角括号
BRACKET_MAP = {
    # 半角括号
    '(': '（',
    ')': '）',
    # 上标括号
    '⁽': '（',
    '⁾': '）',
    # 下标括号
    '₍': '（',
    '₎': '）',
    # 其他可能的括号变体
    '﹙': '（',
    '﹚': '）',
    '❨': '（',
    '❩': '）',
    '⟮': '（',
    '⟯': '）',
}


class DocumentFormatter:
    """公文排版格式化器"""

    def __init__(self):
        """初始化格式化器"""
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """设置页面格式"""
        section = self.doc.sections[0]

        # 设置纸张大小（A4）
        section.page_width = PageSettings.PAGE_WIDTH
        section.page_height = PageSettings.PAGE_HEIGHT

        # 设置页边距
        section.top_margin = PageSettings.MARGIN_TOP
        section.bottom_margin = PageSettings.MARGIN_BOTTOM
        section.left_margin = PageSettings.MARGIN_LEFT
        section.right_margin = PageSettings.MARGIN_RIGHT

    def format_document(self, analysis_result: AnalysisResult) -> Document:
        """
        根据分析结果格式化文档

        Args:
            analysis_result: LLM分析结果

        Returns:
            Document: 格式化后的Word文档
        """
        for element in analysis_result.elements:
            self._add_element(element)

        return self.doc

    def _add_element(self, element: DocumentElement):
        """
        添加文档元素

        Args:
            element: 文档元素
        """
        if not element.content.strip():
            return

        # 获取对应的样式
        font_style = FONT_STYLE_MAP.get(element.element_type, GovDocFonts.BODY)
        para_style = PARAGRAPH_STYLE_MAP.get(element.element_type, GovDocParagraphs.BODY)

        # 对内容进行预处理
        content = element.content

        # 对所有内容进行括号标准化（统一转换为中文全角括号）
        content = self._normalize_brackets(content)

        # 添加段落
        paragraph = self.doc.add_paragraph()

        # 应用段落格式
        self._apply_paragraph_style(paragraph, para_style, element.element_type)

        # 处理混合字体（中文和数字/英文使用不同字体）
        self._add_mixed_font_text(paragraph, content, font_style)

    def _normalize_brackets(self, text: str) -> str:
        """
        标准化括号 - 将各种形式的括号统一转换为中文全角括号，并清理多余空格

        Args:
            text: 原始文本

        Returns:
            str: 括号标准化后的文本
        """
        result = text
        # 替换各种括号为全角括号
        for old_bracket, new_bracket in BRACKET_MAP.items():
            result = result.replace(old_bracket, new_bracket)

        # 清理括号周围的多余空格
        # 左括号后的空格
        result = re.sub(r'（\s+', '（', result)
        # 右括号前的空格
        result = re.sub(r'\s+）', '）', result)

        return result

    def _apply_paragraph_style(self, paragraph, style: ParagraphStyle, element_type: str = None):
        """
        应用段落样式

        Args:
            paragraph: 段落对象
            style: 段落样式
            element_type: 元素类型
        """
        pf = paragraph.paragraph_format

        # 对齐方式
        pf.alignment = style.alignment

        # 首行缩进 - 使用字符单位
        if style.first_line_indent:
            # 通过XML设置首行缩进为2字符
            self._set_first_line_indent_chars(paragraph, 2)

        # 行距（固定值）
        if style.line_spacing:
            pf.line_spacing = style.line_spacing
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        # 段前段后间距
        if style.space_before:
            pf.space_before = style.space_before
        else:
            pf.space_before = Pt(0)

        if style.space_after:
            pf.space_after = style.space_after
        else:
            pf.space_after = Pt(0)

    def _set_first_line_indent_chars(self, paragraph, chars: int):
        """
        设置首行缩进（按字符数）

        Args:
            paragraph: 段落对象
            chars: 缩进的字符数
        """
        # 获取或创建段落属性元素
        pPr = paragraph._p.get_or_add_pPr()

        # 获取或创建缩进元素
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)

        # 设置首行缩进为指定字符数（使用 firstLineChars 属性）
        # 单位是百分之一字符，所以2个字符 = 200
        ind.set(qn('w:firstLineChars'), str(chars * 100))

    def _add_mixed_font_text(self, paragraph, text: str, font_style: FontStyle):
        """
        添加混合字体文本（中文及中文标点用中文字体，仅数字和英文字母用Times New Roman）

        Args:
            paragraph: 段落对象
            text: 文本内容
            font_style: 字体样式
        """
        # 只匹配纯数字和英文字母，其他所有字符（包括中文标点）都用中文字体
        # [a-zA-Z0-9]+ 匹配英文和数字
        # [^a-zA-Z0-9]+ 匹配其他所有字符（中文、标点等）
        pattern = r'([a-zA-Z0-9]+|[^a-zA-Z0-9]+)'
        segments = re.findall(pattern, text)

        for segment in segments:
            if not segment:
                continue

            run = paragraph.add_run(segment)

            # 设置字号
            run.font.size = font_style.size

            # 设置加粗
            run.font.bold = font_style.bold

            # 确保 rPr 元素存在
            rPr = run._element.get_or_add_rPr()

            # 获取或创建 rFonts 元素
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)

            # 判断是否为纯英文数字
            if self._is_ascii_alphanumeric(segment):
                # 数字和英文字母使用 Times New Roman
                rFonts.set(qn('w:ascii'), font_style.name_ascii)
                rFonts.set(qn('w:hAnsi'), font_style.name_ascii)
                rFonts.set(qn('w:eastAsia'), font_style.name)  # 东亚字体也设置
            else:
                # 中文及所有标点符号使用中文字体
                # 同时设置所有字体属性，确保完整覆盖
                rFonts.set(qn('w:ascii'), font_style.name)
                rFonts.set(qn('w:hAnsi'), font_style.name)
                rFonts.set(qn('w:eastAsia'), font_style.name)
                rFonts.set(qn('w:cs'), font_style.name)  # 复杂脚本字体

    def _is_ascii_alphanumeric(self, text: str) -> bool:
        """
        判断文本是否仅包含英文字母和数字

        Args:
            text: 文本

        Returns:
            bool: 是否为纯英文数字
        """
        if not text:
            return False
        return all(c.isascii() and c.isalnum() for c in text)

    def _is_chinese(self, text: str) -> bool:
        """
        判断文本是否主要是中文

        Args:
            text: 文本

        Returns:
            bool: 是否为中文
        """
        if not text:
            return False

        chinese_count = 0
        for char in text:
            if '\u4e00' <= char <= '\u9fff':
                chinese_count += 1

        return chinese_count > len(text) / 2

    def save(self, output_path: str | Path):
        """
        保存文档

        Args:
            output_path: 输出路径
        """
        self.doc.save(str(output_path))


class TitleFormatter:
    """公文标题格式化器（处理多行标题的梯形/菱形排列）"""

    @staticmethod
    def format_title(title: str, max_chars_per_line: int = 22) -> List[str]:
        """
        格式化标题，处理长标题的换行

        Args:
            title: 标题文本
            max_chars_per_line: 每行最大字符数

        Returns:
            List[str]: 分行后的标题列表
        """
        if len(title) <= max_chars_per_line:
            return [title]

        # 尝试在合适的位置断行
        lines = []
        remaining = title

        while remaining:
            if len(remaining) <= max_chars_per_line:
                lines.append(remaining)
                break

            # 寻找断行点（优先在"的"、"关于"等词后断行）
            break_point = TitleFormatter._find_break_point(
                remaining, max_chars_per_line
            )
            lines.append(remaining[:break_point])
            remaining = remaining[break_point:]

        return lines

    @staticmethod
    def _find_break_point(text: str, max_len: int) -> int:
        """
        寻找合适的断行点

        Args:
            text: 文本
            max_len: 最大长度

        Returns:
            int: 断行位置
        """
        if len(text) <= max_len:
            return len(text)

        # 优先在这些词后面断行
        break_after = ["的", "关于", "通知", "报告", "意见", "办法", "规定"]

        for i in range(min(max_len, len(text)) - 1, max_len // 2, -1):
            for word in break_after:
                if text[i - len(word) + 1:i + 1] == word:
                    return i + 1

        # 如果找不到合适的断点，直接在max_len处断开
        return max_len


def format_document(analysis_result: AnalysisResult, output_path: str | Path) -> Path:
    """
    便捷函数：格式化并保存文档

    Args:
        analysis_result: 分析结果
        output_path: 输出路径

    Returns:
        Path: 输出文件路径
    """
    formatter = DocumentFormatter()
    formatter.format_document(analysis_result)
    output_path = Path(output_path)
    formatter.save(output_path)
    return output_path
