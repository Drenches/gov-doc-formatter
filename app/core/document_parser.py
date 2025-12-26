"""
文档解析模块 - 解析Word文档，提取文本内容
"""
import subprocess
import tempfile
import shutil
from pathlib import Path
from typing import List, Dict, Any, Optional
from docx import Document
from docx.text.paragraph import Paragraph
from dataclasses import dataclass, field


@dataclass
class ParsedParagraph:
    """解析后的段落数据"""
    index: int                      # 段落索引
    text: str                       # 段落文本
    original_style: Optional[str]   # 原始样式名称
    is_bold: bool = False           # 是否加粗
    is_empty: bool = False          # 是否为空段落
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class ParsedDocument:
    """解析后的文档数据"""
    filename: str                           # 文件名
    paragraphs: List[ParsedParagraph]       # 段落列表
    full_text: str                          # 完整文本
    paragraph_count: int                    # 段落总数
    metadata: Dict[str, Any] = field(default_factory=dict)


def convert_doc_to_docx(doc_path: Path) -> Path:
    """
    使用 LibreOffice 将 .doc 转换为 .docx

    Args:
        doc_path: .doc 文件路径

    Returns:
        Path: 转换后的 .docx 文件路径

    Raises:
        RuntimeError: 转换失败时抛出
    """
    # 检查 LibreOffice 是否可用
    libreoffice_cmd = None
    for cmd in ['libreoffice', 'soffice', '/usr/bin/libreoffice', '/usr/bin/soffice']:
        if shutil.which(cmd):
            libreoffice_cmd = cmd
            break

    if not libreoffice_cmd:
        raise RuntimeError(
            "未找到 LibreOffice。请安装 LibreOffice:\n"
            "  Ubuntu/Debian: sudo apt install libreoffice\n"
            "  或者将文件另存为 .docx 格式后重新上传"
        )

    # 创建临时目录用于输出
    output_dir = doc_path.parent

    try:
        # 使用 LibreOffice 转换
        result = subprocess.run(
            [
                libreoffice_cmd,
                '--headless',
                '--convert-to', 'docx',
                '--outdir', str(output_dir),
                str(doc_path)
            ],
            capture_output=True,
            text=True,
            timeout=60  # 60秒超时
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")

        # 查找生成的 .docx 文件
        docx_path = doc_path.with_suffix('.docx')
        if not docx_path.exists():
            # 有时文件名可能有细微差异，尝试查找
            possible_files = list(output_dir.glob(f"{doc_path.stem}*.docx"))
            if possible_files:
                docx_path = possible_files[0]
            else:
                raise RuntimeError("转换后未找到 .docx 文件")

        return docx_path

    except subprocess.TimeoutExpired:
        raise RuntimeError("LibreOffice 转换超时")
    except Exception as e:
        raise RuntimeError(f"文档转换失败: {str(e)}")


class DocumentParser:
    """Word文档解析器"""

    def __init__(self, file_path: str | Path):
        """
        初始化解析器

        Args:
            file_path: Word文档路径
        """
        self.file_path = Path(file_path)
        self._converted_file: Optional[Path] = None

        if not self.file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        if self.file_path.suffix.lower() not in ['.docx', '.doc']:
            raise ValueError(f"不支持的文件格式: {self.file_path.suffix}")

        # 如果是 .doc 文件，先转换为 .docx
        if self.file_path.suffix.lower() == '.doc':
            self._converted_file = convert_doc_to_docx(self.file_path)
            self._working_file = self._converted_file
        else:
            self._working_file = self.file_path

    def __del__(self):
        """清理转换的临时文件"""
        if self._converted_file and self._converted_file.exists():
            try:
                self._converted_file.unlink()
            except Exception:
                pass

    def parse(self) -> ParsedDocument:
        """
        解析文档

        Returns:
            ParsedDocument: 解析后的文档数据
        """
        doc = Document(str(self._working_file))
        paragraphs = []
        text_parts = []

        for idx, para in enumerate(doc.paragraphs):
            parsed_para = self._parse_paragraph(idx, para)
            paragraphs.append(parsed_para)
            if parsed_para.text.strip():
                text_parts.append(parsed_para.text)

        return ParsedDocument(
            filename=self.file_path.name,
            paragraphs=paragraphs,
            full_text="\n".join(text_parts),
            paragraph_count=len(paragraphs),
            metadata={
                "file_path": str(self.file_path),
                "core_properties": self._get_core_properties(doc)
            }
        )

    def _parse_paragraph(self, index: int, para: Paragraph) -> ParsedParagraph:
        """
        解析单个段落

        Args:
            index: 段落索引
            para: 段落对象

        Returns:
            ParsedParagraph: 解析后的段落数据
        """
        text = para.text.strip()
        is_empty = len(text) == 0

        # 检查是否加粗
        is_bold = False
        if para.runs:
            # 如果第一个run是加粗的，认为整个段落是加粗的
            is_bold = para.runs[0].bold if para.runs[0].bold is not None else False

        # 获取样式名称
        style_name = None
        if para.style:
            style_name = para.style.name

        return ParsedParagraph(
            index=index,
            text=text,
            original_style=style_name,
            is_bold=is_bold,
            is_empty=is_empty,
            metadata={
                "run_count": len(para.runs),
                "alignment": str(para.alignment) if para.alignment else None
            }
        )

    def _get_core_properties(self, doc: Document) -> Dict[str, Any]:
        """
        获取文档核心属性

        Args:
            doc: 文档对象

        Returns:
            Dict: 核心属性字典
        """
        props = {}
        try:
            core_props = doc.core_properties
            props = {
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "created": str(core_props.created) if core_props.created else None,
                "modified": str(core_props.modified) if core_props.modified else None,
            }
        except Exception:
            pass
        return props

    def get_text_for_llm(self) -> str:
        """
        获取用于发送给LLM的文本格式

        Returns:
            str: 格式化的文本，每段带编号
        """
        parsed = self.parse()
        lines = []

        for para in parsed.paragraphs:
            if not para.is_empty:
                lines.append(f"[{para.index}] {para.text}")

        return "\n".join(lines)

    def cleanup(self):
        """手动清理转换的临时文件"""
        if self._converted_file and self._converted_file.exists():
            try:
                self._converted_file.unlink()
            except Exception:
                pass


def parse_document(file_path: str | Path) -> ParsedDocument:
    """
    便捷函数：解析文档

    Args:
        file_path: 文档路径

    Returns:
        ParsedDocument: 解析后的文档
    """
    parser = DocumentParser(file_path)
    result = parser.parse()
    parser.cleanup()
    return result


def get_document_text(file_path: str | Path) -> str:
    """
    便捷函数：获取文档纯文本

    Args:
        file_path: 文档路径

    Returns:
        str: 文档纯文本
    """
    parser = DocumentParser(file_path)
    result = parser.get_text_for_llm()
    parser.cleanup()
    return result
